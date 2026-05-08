import os
import docx

RAW_LAWS_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "data", "raw_laws")

def generate_mock_html():
    content = """
    <html>
        <head><title>Luật Giao thông Đường bộ 2008</title></head>
        <body>
            <header>Bộ Giao thông Vận tải</header>
            <nav>Home > Luật</nav>
            <h1>Luật Giao thông đường bộ 2008</h1>
            <p>Chương I QUY ĐỊNH CHUNG</p>
            <p>Điều 1. Phạm vi điều chỉnh</p>
            <p>Luật này quy định về quy tắc giao thông đường bộ.</p>
            <p>Điều 2. Đối tượng áp dụng</p>
            <p>1. Tổ chức, cá nhân liên quan đến giao thông đường bộ.</p>
            <p>2. Áp dụng cho cả người nước ngoài.</p>
            <footer>Trang cuối</footer>
        </body>
    </html>
    """
    os.makedirs(RAW_LAWS_DIR, exist_ok=True)
    with open(os.path.join(RAW_LAWS_DIR, "luat_23_2008.html"), "w", encoding="utf-8") as f:
        f.write(content.strip())

def generate_mock_docx():
    os.makedirs(RAW_LAWS_DIR, exist_ok=True)
    doc = docx.Document()
    doc.add_heading('Luật Trật tự, An toàn giao thông đường bộ 2024', 0)
    
    doc.add_paragraph('Chương I QUY ĐỊNH CHUNG')
    doc.add_paragraph('Điều 1. Phạm vi điều chỉnh')
    doc.add_paragraph('1. Luật này quy định về quy tắc trật tự, an toàn giao thông đường bộ.')
    doc.add_paragraph('Điều 2. Đối tượng áp dụng')
    doc.add_paragraph('Áp dụng đối với tổ chức, cá nhân.')
    
    doc.save(os.path.join(RAW_LAWS_DIR, "luat_36_2024.docx"))

def generate_mock_data():
    generate_mock_html()
    try:
        generate_mock_docx()
        print("Đã tạo mock data HTML và DOCX thành công.")
    except Exception as e:
        print(f"Lỗi khi tạo file docx: {e}")

if __name__ == "__main__":
    generate_mock_data()
